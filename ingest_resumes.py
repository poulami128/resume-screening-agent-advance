# ingest_resumes.py
import io
from typing import List, Dict
import os

def load_resumes_from_folder(folder: str) -> List[Dict]:
    results = []
    if not os.path.isdir(folder):
        return results
    for fn in os.listdir(folder):
        path = os.path.join(folder, fn)
        if not os.path.isfile(path):
            continue
        try:
            text = _extract_text_from_path(path)
            results.append({"name": fn, "text": text})
        except Exception:
            results.append({"name": fn, "text": ""})
    return results

def extract_text_from_uploaded_file(uploaded_file) -> str:
    # uploaded_file is a Streamlit UploadedFile object
    name = uploaded_file.name.lower()
    content = uploaded_file.read()
    # for pdf
    if name.endswith(".pdf"):
        try:
            from pdfminer.high_level import extract_text
            # pdfminer can accept a file-like object
            with io.BytesIO(content) as b:
                return extract_text(b) or ""
        except Exception:
            return ""
    # for docx
    if name.endswith(".docx"):
        try:
            import docx
            from io import BytesIO
            doc = docx.Document(BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs]
            return "\n".join(paragraphs)
        except Exception:
            return ""
    # for txt
    try:
        return content.decode("utf-8", errors="ignore")
    except Exception:
        return ""

def _extract_text_from_path(path: str) -> str:
    name = path.lower()
    if path.endswith(".pdf"):
        from pdfminer.high_level import extract_text
        return extract_text(path) or ""
    elif path.endswith(".docx"):
        import docx
        doc = docx.Document(path)
        return "\n".join(p.text for p in doc.paragraphs)
    else:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
